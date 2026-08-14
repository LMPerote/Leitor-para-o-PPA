"""Testes do extrator.

Os documentos de teste são gerados em tempo de execução com python-docx,
cobrindo as situações que aparecem no acervo real:

* os dois tipos de ficha (Indicador de Compromisso e Iniciativa);
* layout vertical  (rótulo em uma linha, valor na linha de baixo);
* layout horizontal (rótulo à esquerda, valor à direita);
* layout em parágrafos com "Rótulo: valor" na mesma linha;
* arquivo corrompido e arquivo de tipo não reconhecido.

Execute com:  pytest -q
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx
import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from extrator import ler_documento, montar_dataframe, processar_lote  # noqa: E402
from extrator.modelos import MODELOS_POR_CODIGO  # noqa: E402
from extrator.parser import Extrator, classificar, eh_rotulo  # noqa: E402
from extrator.pipeline import listar_documentos  # noqa: E402
from extrator.planilha import LIMITE_CELULA, salvar_excel, sanitizar  # noqa: E402
from extrator.texto import chave, esta_marcado  # noqa: E402

INDICADOR = MODELOS_POR_CODIGO["INDICADOR"]
INICIATIVA = MODELOS_POR_CODIGO["INICIATIVA"]


# ---------------------------------------------------------------------------
# Auxiliares de construção de documentos
# ---------------------------------------------------------------------------
def _tabela_vertical(documento, linhas: list[str]) -> None:
    """Tabela de uma coluna: rótulo, valor, rótulo, valor..."""
    tabela = documento.add_table(rows=len(linhas), cols=1)
    for indice, texto in enumerate(linhas):
        tabela.cell(indice, 0).text = texto


def _tabela_horizontal(documento, pares: list[tuple[str, str]]) -> None:
    """Tabela de duas colunas: rótulo | valor."""
    tabela = documento.add_table(rows=len(pares), cols=2)
    for indice, (rotulo, valor) in enumerate(pares):
        tabela.cell(indice, 0).text = rotulo
        tabela.cell(indice, 1).text = valor


def _tabela_grade(documento, linhas: list[list[str]]) -> None:
    """Tabela com cabeçalhos na primeira linha e valores nas seguintes."""
    colunas = max(len(linha) for linha in linhas)
    tabela = documento.add_table(rows=len(linhas), cols=colunas)
    for l, linha in enumerate(linhas):
        for c, texto in enumerate(linha):
            tabela.cell(l, c).text = texto


def criar_ficha_indicador(destino: Path) -> Path:
    """Ficha de Indicador no layout padrão (misto vertical/horizontal)."""
    documento = docx.Document()

    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_vertical(
        documento,
        [
            "Eixo",
            "EIXO DE TESTE",
            "Programa",
            "Programa de Teste",
            "Compromisso",
            "Compromisso de teste",
            "Problema(s) vinculado(s) ao Compromisso",
            "Problema de teste",
            "Causa(s) Crítica(s)",
            "Causa 1.\nCausa 2.\nCausa 3.",
            "ATRIBUTOSDO INDICADOR DE COMPROMISSO",  # erro de espaço proposital
        ],
    )
    _tabela_vertical(
        documento,
        [
            "Descrição",
            "Descrição do indicador",
            "Fórmula de Cálculo",
            "Somatório de X",
            "Memória de Cálculo",
            "Memória detalhada",
        ],
    )
    _tabela_grade(
        documento,
        [
            ["Unidade de medida", "Valor de referência", "Ano de referência", "Valor da meta"],
            ["Unidade", "6", "2023", "40"],
            ["Periodicidade da apuração", "Polaridade", "Classificação", ""],
            ["Semestral", "Positiva", "Produto", ""],
        ],
    )
    _tabela_vertical(
        documento,
        ["Fonte", "Órgão XPTO", "Meios de verificação", "Relatórios administrativos"],
    )
    _tabela_grade(documento, [["Sigla do Órgão", "UO", "USP"], ["SJDH", "APG", "O16 GASEC"]])

    _tabela_grade(
        documento,
        [
            ["DESAGREGAÇÃO TERRITORIAL", "", "", ""],
            ["Estado", "x", "Território de Identidade", ""],
            ["Fórmula de cálculo Territorial", "", "Unidade de Medida", ""],
            ["Fórmula territorial", "", "Unidade", ""],
            ["Memória de Cálculo", "", "", ""],
            ["Sem memória", "", "", ""],
            [
                "Território de Identidade",
                "Memória de Cálculo Territorial",
                "Meta Territorial",
                "",
            ],
            ["Metropolitano", "Soma regional", "10", ""],
            ["Recôncavo", "Soma regional", "5", ""],
            ["Outras possibilidades de Regionalização", "", "", ""],
            ["Não se aplica", "", "", ""],
        ],
    )

    _tabela_vertical(documento, ["INFORMAÇÕES COMPLEMENTARES"])
    _tabela_vertical(
        documento,
        [
            "Objetivo/ Interpretação e uso",
            "Mede a evolução de X",
            "Limitações do Indicador",
            "Não se aplica",
            "Fragilidades para apuração e ações em curso para superação",
            "Nenhuma",
        ],
    )
    _tabela_horizontal(
        documento,
        [
            ("Limitações para definição do valor da meta", ""),
            ("Operacionais", "Limitação operacional"),
            ("Orçamentárias/Financeiras", "Limitação orçamentária"),
            ("Institucionais ou políticas", "Limitação institucional"),
        ],
    )
    _tabela_vertical(documento, ["Possibilidade de desagregação populacional", "Sim, por sexo"])
    _tabela_grade(
        documento,
        [
            ["Programas Especiais", "", ""],
            ["Nome do Programa", "Memória de Cálculo", "Meta"],
            ["Programa Especial A", "Soma", "12"],
        ],
    )

    documento.save(destino)
    return destino


def criar_ficha_iniciativa(destino: Path) -> Path:
    """Ficha de Iniciativa no layout padrão."""
    documento = docx.Document()

    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Eixo",
            "EIXO DA INICIATIVA",
            "Programa",
            "Programa da Iniciativa",
            "Compromisso",
            "Compromisso da iniciativa",
            "Problema(s) vinculado(s) ao Compromisso",
            "Problema da iniciativa",
            "Causa(s) Crítica(s)",
            "Causa A.\nCausa B.",
            "Ação(ões) Crítica(s)",
            "Ação 1.\nAção 2.",
        ],
    )
    _tabela_grade(
        documento,
        [
            [
                "Proposta(s) de Escuta Social Associada(s) ao Compromisso",
                "Status para atendimento",
                "Justificativa(s)",
            ],
            ["Proposta A", "Atendida parcialmente", "-"],
            ["Proposta B", "Não atendida", "Sem orçamento"],
        ],
    )

    _tabela_vertical(documento, ["ATRIBUTOS DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Descrição",
            "Descrição da iniciativa",
            "Entrega(s) Vinculada(s)",
            "Entrega 1\nEntrega 2",
            "Responsável pela Iniciativa",
        ],
    )
    _tabela_grade(
        documento,
        [
            ["Sigla do Órgão", "UO", "USP", "Órgãos Parceiros"],
            ["SJDH", "APG", "GASEC", "Parceiro A, Parceiro B"],
        ],
    )
    _tabela_grade(
        documento,
        [
            ["Fatores Críticos de Contexto da Iniciativa", ""],
            ["Operacionais", "Fator operacional"],
            ["Orçamentários/Financeiros", "Fator orçamentário"],
            ["Institucionais ou políticos", "Fator institucional"],
        ],
    )
    _tabela_grade(
        documento,
        [
            [
                "Estimativa dos recursos orçamentários/financeiros disponíveis para Iniciativa",
                "",
                "",
            ],
            ["Código da Fonte", "Nome da Fonte", "Montante em R$"],
            ["100", "Tesouro Estadual", "R$ 2.400.000,00"],
            ["Total dos Recursos", "", "R$ 2.400.000,00"],
        ],
    )
    _tabela_vertical(
        documento,
        [
            "Indicador de Compromisso Vinculado",
            "Número de campanhas realizadas",
            "Indicadores de Compromisso Sensibilizados",
            "Indicador sensibilizado A",
        ],
    )
    _tabela_grade(
        documento,
        [
            [
                "Etapa II – Programação Orçamentária - Ações orçamentárias vinculadas à Iniciativa",
                "",
                "",
                "",
            ],
            ["Órgão", "UO", "Código da Ação Orçamentária", "Nome da Ação Orçamentária"],
            ["SJDH", "APG", "4170", "Realização de ações de promoção"],
        ],
    )
    _tabela_grade(
        documento,
        [
            [
                "Etapa II – Programação Orçamentária – Produtos das ações orçamentárias "
                "vinculadas à Iniciativa",
                "",
                "",
                "",
                "",
            ],
            [
                "Órgão",
                "UO",
                "Código da Ação Orçamentária",
                "Código do Produto da Ação Orçamentária",
                "Nome Produto da Ação Orçamentária",
            ],
            ["SJDH", "APG", "4170", "P1", "Ações de promoção realizadas"],
        ],
    )

    documento.save(destino)
    return destino


def criar_ficha_em_paragrafos(destino: Path) -> Path:
    """Variação fora do padrão: texto corrido com "Rótulo: valor"."""
    documento = docx.Document()
    documento.add_paragraph("VÍNCULO DO INDICADOR DE COMPROMISSO")
    documento.add_paragraph("Eixo: EIXO EM PARÁGRAFO")
    documento.add_paragraph("Programa")
    documento.add_paragraph("Programa em parágrafo")
    documento.add_paragraph("ATRIBUTOS DO INDICADOR DE COMPROMISSO")
    documento.add_paragraph("Descrição: Indicador descrito em parágrafo")
    documento.add_paragraph("Fonte: Órgão do parágrafo")
    documento.save(destino)
    return destino


def criar_documento_desconhecido(destino: Path) -> Path:
    """Um .docx válido que não é nenhuma das duas fichas."""
    documento = docx.Document()
    documento.add_paragraph("MEMORANDO INTERNO")
    documento.add_paragraph("Este documento não é uma ficha do PPA.")
    documento.save(destino)
    return destino


@pytest.fixture(scope="module")
def pasta(tmp_path_factory) -> Path:
    destino = tmp_path_factory.mktemp("fichas")
    criar_ficha_indicador(destino / "ficha_indicador.docx")
    criar_ficha_iniciativa(destino / "ficha_iniciativa.docx")
    criar_ficha_em_paragrafos(destino / "ficha_paragrafos.docx")
    criar_documento_desconhecido(destino / "outro_documento.docx")
    (destino / "corrompido.docx").write_bytes(b"isto nao e um docx valido")
    (destino / "~$temporario.docx").write_bytes(b"lixo")
    (destino / "ignorar.txt").write_text("nao e docx")
    return destino


@pytest.fixture(scope="module")
def indicador(pasta: Path) -> dict[str, str]:
    documento = ler_documento(str(pasta / "ficha_indicador.docx"))
    return Extrator(INDICADOR).extrair(documento).valores


@pytest.fixture(scope="module")
def iniciativa(pasta: Path) -> dict[str, str]:
    documento = ler_documento(str(pasta / "ficha_iniciativa.docx"))
    return Extrator(INICIATIVA).extrair(documento).valores


# ---------------------------------------------------------------------------
# Normalização de texto
# ---------------------------------------------------------------------------
def test_chave_ignora_acentos_espacos_e_pontuacao():
    assert chave("Fórmula de Cálculo") == "formuladecalculo"
    assert chave("Indicador(es) doPrograma Sensibilizado(s)") == chave(
        "Indicador(es) do Programa Sensibilizado(s)"
    )


def test_reconhecimento_de_rotulos():
    assert eh_rotulo("Valor de referência")
    assert eh_rotulo("ATRIBUTOSDO INDICADOR DE COMPROMISSO")
    assert eh_rotulo("Proposta(s) de Escuta Social Associada(s) ao Compromisso")
    assert not eh_rotulo("Semestral")
    assert not eh_rotulo("40")


def test_marcacao_de_caixa_de_selecao():
    assert esta_marcado("x") and esta_marcado("☒")
    assert not esta_marcado("") and not esta_marcado("☐")


# ---------------------------------------------------------------------------
# Classificação automática
# ---------------------------------------------------------------------------
def test_classificacao_por_secao_de_vinculo(pasta: Path):
    def tipo(nome: str):
        return classificar(ler_documento(str(pasta / nome)))

    assert tipo("ficha_indicador.docx") is INDICADOR
    assert tipo("ficha_iniciativa.docx") is INICIATIVA
    assert tipo("ficha_paragrafos.docx") is INDICADOR  # vínculo em parágrafo
    assert tipo("outro_documento.docx") is None


# ---------------------------------------------------------------------------
# Campos da ficha de INDICADOR
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    ("coluna", "esperado"),
    [
        ("Eixo", "EIXO DE TESTE"),
        ("Programa", "Programa de Teste"),
        ("Compromisso", "Compromisso de teste"),
        ("Problemas_Vinculados", "Problema de teste"),
        ("Atributos_Descricao", "Descrição do indicador"),
        ("Formula_Calculo", "Somatório de X"),
        ("Memoria_Calculo", "Memória detalhada"),
        ("Unidade_Medida", "Unidade"),
        ("Valor_Referencia", "6"),
        ("Ano_Referencia", "2023"),
        ("Valor_Meta", "40"),
        ("Periodicidade_Apuracao", "Semestral"),
        ("Polaridade", "Positiva"),
        ("Classificacao", "Produto"),
        ("Fonte", "Órgão XPTO"),
        ("Meios_Verificacao", "Relatórios administrativos"),
        ("Responsavel_Sigla_Orgao", "SJDH"),
        ("Responsavel_UO", "APG"),
        ("Responsavel_USP", "O16 GASEC"),
        ("Info_Complementares_Objetivo", "Mede a evolução de X"),
        ("Limitacoes_Indicador", "Não se aplica"),
        ("Fragilidades_Apuracao", "Nenhuma"),
        ("Limitacoes_Meta_Operacionais", "Limitação operacional"),
        ("Limitacoes_Meta_Orcamentarias", "Limitação orçamentária"),
        ("Limitacoes_Meta_Institucionais", "Limitação institucional"),
        ("Possibilidade_Desagregacao_Populacional", "Sim, por sexo"),
    ],
)
def test_indicador_campos_simples(indicador, coluna, esperado):
    assert indicador[coluna] == esperado


def test_indicador_lista_permanece_em_uma_unica_celula(indicador):
    assert indicador["Causas_Criticas"] == "Causa 1.\nCausa 2.\nCausa 3."


def test_indicador_desagregacao_territorial_consolidada(indicador):
    territorial = indicador["Desagregacao_Territorial"]
    assert "Estado: Sim" in territorial
    assert "Território de Identidade: Não" in territorial
    assert "Fórmula de cálculo Territorial: Fórmula territorial" in territorial
    assert "Unidade de Medida: Unidade" in territorial
    assert "Memória de Cálculo: Sem memória" in territorial
    assert "Territórios de Identidade: Metropolitano\nRecôncavo" in territorial
    assert "Meta Territorial: 10\n5" in territorial
    assert "Outras possibilidades de Regionalização: Não se aplica" in territorial


def test_indicador_programas_especiais(indicador):
    assert indicador["Programas_Especiais"] == (
        "Nome do Programa: Programa Especial A | Memória de Cálculo: Soma | Meta: 12"
    )


# ---------------------------------------------------------------------------
# Campos da ficha de INICIATIVA
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    ("coluna", "esperado"),
    [
        ("Eixo", "EIXO DA INICIATIVA"),
        ("Programa", "Programa da Iniciativa"),
        ("Compromisso", "Compromisso da iniciativa"),
        ("Problemas_Vinculados", "Problema da iniciativa"),
        ("Causas_Criticas", "Causa A.\nCausa B."),
        ("Acoes_Criticas", "Ação 1.\nAção 2."),
        ("Atributos_Descricao", "Descrição da iniciativa"),
        ("Entregas_Vinculadas", "Entrega 1\nEntrega 2"),
        ("Responsavel_Sigla_Orgao", "SJDH"),
        ("Responsavel_UO", "APG"),
        ("Responsavel_USP", "GASEC"),
        ("Orgaos_Parceiros", "Parceiro A, Parceiro B"),
        ("Fatores_Criticos_Operacionais", "Fator operacional"),
        ("Fatores_Criticos_Orcamentarios", "Fator orçamentário"),
        ("Fatores_Criticos_Institucionais", "Fator institucional"),
        ("Indicador_Compromisso_Vinculado", "Número de campanhas realizadas"),
        ("Indicadores_Sensibilizados", "Indicador sensibilizado A"),
    ],
)
def test_iniciativa_campos_simples(iniciativa, coluna, esperado):
    assert iniciativa[coluna] == esperado


def test_iniciativa_propostas_de_escuta_social(iniciativa):
    linhas = iniciativa["Propostas_Escuta_Social"].split("\n")
    assert len(linhas) == 2
    assert linhas[0] == (
        "Proposta(s) de Escuta Social Associada(s) ao Compromisso: Proposta A | "
        "Status para atendimento: Atendida parcialmente | Justificativa(s): -"
    )
    assert "Justificativa(s): Sem orçamento" in linhas[1]


def test_iniciativa_recursos_orcamentarios_incluem_o_total(iniciativa):
    recursos = iniciativa["Recursos_Orcamentarios"].split("\n")
    assert recursos[0] == (
        "Código da Fonte: 100 | Nome da Fonte: Tesouro Estadual | "
        "Montante em R$: R$ 2.400.000,00"
    )
    assert recursos[-1] == "Total dos Recursos: R$ 2.400.000,00"


def test_iniciativa_acoes_e_produtos_orcamentarios(iniciativa):
    assert iniciativa["Acoes_Orcamentarias_Vinculadas"] == (
        "Órgão: SJDH | UO: APG | Código da Ação Orçamentária: 4170 | "
        "Nome da Ação Orçamentária: Realização de ações de promoção"
    )
    produtos = iniciativa["Produtos_Acoes_Orcamentarias"]
    assert "Código do Produto da Ação Orçamentária: P1" in produtos
    assert "Nome Produto da Ação Orçamentária: Ações de promoção realizadas" in produtos


def test_separador_configuravel(pasta: Path):
    documento = ler_documento(str(pasta / "ficha_iniciativa.docx"))
    valores = Extrator(INICIATIVA, separador="; ").extrair(documento).valores
    assert valores["Acoes_Criticas"] == "Ação 1.; Ação 2."


# ---------------------------------------------------------------------------
# Layouts fora do padrão
# ---------------------------------------------------------------------------
def test_layout_em_paragrafos(pasta: Path):
    documento = ler_documento(str(pasta / "ficha_paragrafos.docx"))
    valores = Extrator(INDICADOR).extrair(documento).valores
    assert valores["Eixo"] == "EIXO EM PARÁGRAFO"  # "Rótulo: valor" na mesma linha
    assert valores["Programa"] == "Programa em parágrafo"  # valor no parágrafo seguinte
    assert valores["Atributos_Descricao"] == "Indicador descrito em parágrafo"
    assert valores["Fonte"] == "Órgão do parágrafo"


def test_campos_ausentes_sao_reportados(pasta: Path):
    documento = ler_documento(str(pasta / "ficha_paragrafos.docx"))
    resultado = Extrator(INDICADOR).extrair(documento)
    assert "Polaridade" in resultado.nao_encontrados
    assert "Eixo" not in resultado.nao_encontrados


# ---------------------------------------------------------------------------
# Lote, resiliência e planilhas
# ---------------------------------------------------------------------------
def test_listagem_ignora_temporarios_e_outras_extensoes(pasta: Path):
    nomes = {caminho.name for caminho in listar_documentos(pasta)}
    assert nomes == {
        "ficha_indicador.docx",
        "ficha_iniciativa.docx",
        "ficha_paragrafos.docx",
        "outro_documento.docx",
        "corrompido.docx",
    }


def test_lote_separa_por_tipo_e_nao_para_em_erros(pasta: Path):
    arquivos = listar_documentos(pasta)
    registros, estatisticas = processar_lote(arquivos, pasta)

    assert estatisticas.total == len(arquivos)
    assert estatisticas.processados["INDICADOR"] == 2  # tabelas + parágrafos
    assert estatisticas.processados["INICIATIVA"] == 1
    # Arquivos processados e linhas geradas são coisas diferentes: a ficha de
    # indicador em tabelas traz 3 causas críticas (3 linhas) e a de iniciativa,
    # 2 causas e 2 ações (2 linhas).
    assert len(registros["INDICADOR"]) == 4
    assert len(registros["INICIATIVA"]) == 2
    assert estatisticas.linhas["INDICADOR"] == 4
    assert estatisticas.linhas["INICIATIVA"] == 2

    # Corrompido e "não é ficha do PPA" ficam de fora, mas são reportados.
    ignorados = {
        ocorrencia.arquivo: ocorrencia.motivo for ocorrencia in estatisticas.ignorados
    }
    assert set(ignorados) == {"corrompido.docx", "outro_documento.docx"}
    assert "erro de leitura" in ignorados["corrompido.docx"]
    assert "tipo não reconhecido" in ignorados["outro_documento.docx"]


def test_planilhas_geradas_tem_uma_linha_por_causa_critica(pasta: Path, tmp_path: Path):
    import openpyxl

    arquivos = listar_documentos(pasta)
    registros, _ = processar_lote(arquivos, pasta)

    for modelo, esperado in ((INDICADOR, 4), (INICIATIVA, 2)):
        destino = tmp_path / modelo.arquivo_saida
        salvar_excel(montar_dataframe(registros[modelo.codigo], modelo), destino, modelo)

        planilha = openpyxl.load_workbook(destino)[destino.stem]
        assert planilha.max_row == esperado + 1  # + cabeçalho
        assert planilha.max_column == len(modelo.colunas())
        assert planilha.cell(row=1, column=1).value == "Nome_Arquivo"
        assert planilha.cell(row=1, column=1).font.bold
        assert planilha.cell(row=2, column=1).alignment.wrap_text
        assert planilha.cell(row=2, column=1).alignment.vertical == "top"
        assert planilha.freeze_panes == "A2"  # só o cabeçalho; nenhuma coluna fixa


def test_planilha_vazia_ainda_tem_cabecalho(tmp_path: Path):
    import openpyxl

    destino = tmp_path / INICIATIVA.arquivo_saida
    salvar_excel(montar_dataframe([], INICIATIVA), destino, INICIATIVA)

    planilha = openpyxl.load_workbook(destino)[destino.stem]
    assert planilha.max_row == 1
    assert [celula.value for celula in planilha[1]] == INICIATIVA.colunas()


def test_sanitizacao_para_o_excel():
    assert sanitizar(None) == ""
    assert sanitizar("=1+1").startswith("'")  # não vira fórmula
    assert sanitizar("a\x07b") == "ab"  # caractere de controle removido
    assert len(sanitizar("x" * (LIMITE_CELULA + 500))) <= LIMITE_CELULA


# ---------------------------------------------------------------------------
# Fluxo completo (compartilhado pela linha de comando e pela janela gráfica)
# ---------------------------------------------------------------------------
def test_executar_grava_as_duas_planilhas(pasta: Path, tmp_path: Path):
    from extrator.aplicacao import Opcoes, executar

    saida = tmp_path / "planilhas"
    execucao = executar(Opcoes(entrada=pasta, saida=saida))

    assert (saida / "Indicadores.xlsx").is_file()
    assert (saida / "Iniciativas.xlsx").is_file()
    assert execucao.estatisticas.processados == {"INDICADOR": 2, "INICIATIVA": 1}
    assert execucao.destinos["INDICADOR"] == (saida / "Indicadores.xlsx").resolve()


def test_executar_respeita_limite_e_progresso(pasta: Path, tmp_path: Path):
    from extrator.aplicacao import Opcoes, executar

    class Contador:
        def __init__(self, total):
            self.total, self.passos = total, 0

        def update(self, quantidade=1):
            self.passos += quantidade

    criados: list[Contador] = []

    def criar(total):
        criados.append(Contador(total))
        return criados[-1]

    execucao = executar(
        Opcoes(entrada=pasta, saida=tmp_path / "p", limite=2), criar_progresso=criar
    )
    assert len(execucao.arquivos) == 2
    assert criados[0].total == 2 and criados[0].passos == 2


def test_executar_sem_documentos_avisa(tmp_path: Path):
    from extrator.aplicacao import Opcoes, PastaSemDocumentos, executar

    vazia = tmp_path / "vazia"
    vazia.mkdir()
    with pytest.raises(PastaSemDocumentos):
        executar(Opcoes(entrada=vazia, saida=tmp_path / "p"))


def test_relatorio_traz_contagem_por_tipo_e_ignorados(pasta: Path, tmp_path: Path):
    from extrator.aplicacao import Opcoes, executar
    from extrator.relatorio import montar_relatorio

    execucao = executar(Opcoes(entrada=pasta, saida=tmp_path / "p"))
    texto = montar_relatorio(execucao.estatisticas, execucao.destinos, 1.23)

    assert "2 arquivos de Indicador de Compromisso processados." in texto
    assert "1 arquivos de Iniciativa processados." in texto
    assert "2 arquivos ignorados ou com erro." in texto
    assert "corrompido.docx" in texto and "outro_documento.docx" in texto
    assert "1.2s" in texto


def test_apelidos_de_separador():
    from extrator.aplicacao import resolver_separador

    assert resolver_separador("quebra") == "\n"
    assert resolver_separador("ponto-virgula") == "; "
    assert resolver_separador(" / ") == " / "  # texto literal passa direto


# ---------------------------------------------------------------------------
# Regra de negócio: UM problema por linha
# ---------------------------------------------------------------------------
def _ficha_com_dois_problemas(destino: Path) -> Path:
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_vertical(
        documento,
        [
            "Eixo",
            "EIXO COM DOIS PROBLEMAS",
            "Problema(s) vinculado(s) ao Compromisso",
            "Primeiro problema.\nSegundo problema.",
            "Causa(s) Crítica(s)",
            "Causa única.",
            "ATRIBUTOS DO INDICADOR DE COMPROMISSO",
        ],
    )
    _tabela_vertical(documento, ["Descrição", "Indicador de teste"])
    documento.save(destino)
    return destino


def test_ficha_com_dois_problemas_gera_duas_linhas(tmp_path: Path):
    pasta = tmp_path / "dois"
    pasta.mkdir()
    _ficha_com_dois_problemas(pasta / "ficha.docx")

    registros, estatisticas = processar_lote(listar_documentos(pasta), pasta)
    linhas = registros["INDICADOR"]

    assert [linha["Problemas_Vinculados"] for linha in linhas] == [
        "Primeiro problema.",
        "Segundo problema.",
    ]
    # As demais colunas se repetem em todas as linhas da mesma ficha.
    for coluna in ("Nome_Arquivo", "Eixo", "Causas_Criticas", "Atributos_Descricao"):
        assert linhas[0][coluna] == linhas[1][coluna]
    assert linhas[0]["Eixo"] == "EIXO COM DOIS PROBLEMAS"

    # O arquivo continua contando como UM arquivo processado, com duas linhas.
    assert estatisticas.processados["INDICADOR"] == 1
    assert estatisticas.linhas["INDICADOR"] == 2


def test_iniciativa_gera_uma_linha_por_causa_acao_e_entrega(pasta: Path):
    """A ficha tem 1 problema, 2 causas, 2 ações e 2 entregas -> 2 linhas."""
    registros, _ = processar_lote(listar_documentos(pasta), pasta)
    linhas = registros["INICIATIVA"]
    assert len(linhas) == 2
    assert [linha["Causas_Criticas"] for linha in linhas] == ["Causa A.", "Causa B."]
    assert [linha["Acoes_Criticas"] for linha in linhas] == ["Ação 1.", "Ação 2."]
    # A entrega também é dividida: uma por linha, e não empilhada na célula.
    assert [linha["Entregas_Vinculadas"] for linha in linhas] == ["Entrega 1", "Entrega 2"]
    # O problema é um só: se repete nas duas linhas.
    assert {linha["Problemas_Vinculados"] for linha in linhas} == {"Problema da iniciativa"}
    # As demais colunas acompanham a ficha inteira.
    assert {linha["Atributos_Descricao"] for linha in linhas} == {"Descrição da iniciativa"}
    assert {linha["Nome_Arquivo"] for linha in linhas} == {"ficha_iniciativa.docx"}


def test_expansao_respeita_o_separador_escolhido():
    from extrator.pipeline import expandir_em_linhas

    registro = {"Nome_Arquivo": "f.docx", "Problemas_Vinculados": "A; B; C"}
    linhas = expandir_em_linhas(registro, "; ")
    assert [linha["Problemas_Vinculados"] for linha in linhas] == ["A", "B", "C"]
    assert all(linha["Nome_Arquivo"] == "f.docx" for linha in linhas)


def test_ficha_sem_problema_nao_some_da_planilha():
    from extrator.pipeline import expandir_em_linhas

    registro = {"Nome_Arquivo": "f.docx", "Problemas_Vinculados": ""}
    assert expandir_em_linhas(registro, "\n") == [registro]


def test_problema_causa_e_acao_sao_pareados_pela_ordem():
    """Cada linha traz a trinca que se corresponde dentro da mesma ficha."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1\nP2\nP3",
        "Causas_Criticas": "C1\nC2\nC3",
        "Acoes_Criticas": "A1\nA2\nA3",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert [
        (l["Problemas_Vinculados"], l["Causas_Criticas"], l["Acoes_Criticas"]) for l in linhas
    ] == [("P1", "C1", "A1"), ("P2", "C2", "A2"), ("P3", "C3", "A3")]


def test_lista_menor_repete_o_ultimo_item():
    """4 problemas, 1 causa e 3 ações viram 4 linhas, como na planilha de vínculo."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1\nP2\nP3\nP4",
        "Causas_Criticas": "C1",
        "Acoes_Criticas": "A1\nA2\nA3",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert [
        (l["Problemas_Vinculados"], l["Causas_Criticas"], l["Acoes_Criticas"]) for l in linhas
    ] == [("P1", "C1", "A1"), ("P2", "C1", "A2"), ("P3", "C1", "A3"), ("P4", "C1", "A3")]


def test_mais_causas_que_problemas_repete_o_problema():
    """Um problema que vale para várias causas aparece repetido, sem deduplicar."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1",
        "Causas_Criticas": "C1\nC2\nC3\nC4",
        "Acoes_Criticas": "A1\nA2\nA3\nA4",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert len(linhas) == 4
    assert all(linha["Problemas_Vinculados"] == "P1" for linha in linhas)
    assert [linha["Causas_Criticas"] for linha in linhas] == ["C1", "C2", "C3", "C4"]


def test_expansao_nao_cria_coluna_que_o_modelo_nao_tem():
    """A ficha de Indicador não tem ações críticas; a coluna não é inventada."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1",
        "Causas_Criticas": "C1\nC2",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert len(linhas) == 2
    assert all("Acoes_Criticas" not in linha for linha in linhas)


def test_entrega_tambem_vira_uma_linha_por_item():
    """A entrega entra na mesma regra: 3 entregas geram 3 linhas."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1",
        "Causas_Criticas": "C1",
        "Acoes_Criticas": "A1",
        "Entregas_Vinculadas": "E1\nE2\nE3",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert [linha["Entregas_Vinculadas"] for linha in linhas] == ["E1", "E2", "E3"]
    # O trio se repete nas três linhas, sem agrupar nem deduplicar.
    assert all(
        (l["Problemas_Vinculados"], l["Causas_Criticas"], l["Acoes_Criticas"])
        == ("P1", "C1", "A1")
        for l in linhas
    )


def test_itens_digitados_na_mesma_linha_sao_divididos():
    """Lista escrita seguida ("P1 ...; P2 ...;") também vira uma linha por item."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1 Violência contra jovens; P2 Racismo institucional;",
        "Entregas_Vinculadas": "Carnaval; Festas populares",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert [linha["Problemas_Vinculados"] for linha in linhas] == [
        "P1 Violência contra jovens",
        "P2 Racismo institucional",
    ]
    assert [linha["Entregas_Vinculadas"] for linha in linhas] == [
        "Carnaval",
        "Festas populares",
    ]


def test_item_repetido_ocupa_a_propria_linha():
    """Repetição é preservada: cada ocorrência vale uma linha."""
    from extrator.pipeline import expandir_em_linhas

    registro = {
        "Nome_Arquivo": "f.docx",
        "Problemas_Vinculados": "P1\nP1\nP2",
        "Causas_Criticas": "C1\nC1\nC1",
    }
    linhas = expandir_em_linhas(registro, "\n")
    assert [linha["Problemas_Vinculados"] for linha in linhas] == ["P1", "P1", "P2"]
    assert [linha["Causas_Criticas"] for linha in linhas] == ["C1", "C1", "C1"]


def test_divisao_normaliza_item_unico():
    """Um item só continua gerando uma linha, sem o separador sobrando."""
    from extrator.pipeline import expandir_em_linhas

    registro = {"Nome_Arquivo": "f.docx", "Problemas_Vinculados": "P1;"}
    linhas = expandir_em_linhas(registro, "\n")
    assert len(linhas) == 1
    assert linhas[0]["Problemas_Vinculados"] == "P1"


# ---------------------------------------------------------------------------
# Recuperação de pacotes .docx parcialmente corrompidos
# ---------------------------------------------------------------------------
def _corromper_imagem(origem: Path, destino: Path) -> Path:
    """Reescreve o .docx com o CRC de uma imagem inválido.

    Reproduz o erro real "Bad CRC-32 for file 'word/media/imageN.png'", que
    aparece em fichas gravadas pelo Word com a imagem danificada.
    """
    import zipfile

    with zipfile.ZipFile(origem) as entrada, zipfile.ZipFile(destino, "w") as saida:
        for item in entrada.infolist():
            dados = entrada.read(item.filename)
            if item.filename.startswith("word/media/"):
                info = zipfile.ZipInfo(item.filename, date_time=item.date_time)
                info.compress_type = zipfile.ZIP_STORED
                saida.writestr(info, dados)
            else:
                saida.writestr(item, dados)

    # Vira um byte no meio da imagem gravada sem compressão: o conteúdo deixa
    # de bater com o CRC registrado no cabeçalho.
    bruto = bytearray(destino.read_bytes())
    marca = bruto.find(b"\x89PNG")
    assert marca > 0, "imagem não encontrada no pacote de teste"
    bruto[marca + 20] ^= 0xFF
    destino.write_bytes(bytes(bruto))
    return destino


def test_docx_com_imagem_corrompida_ainda_e_lido(tmp_path: Path):
    import zipfile

    original = criar_ficha_indicador(tmp_path / "com_imagem.docx")
    documento = docx.Document(original)
    documento.add_picture(str(_png_de_teste(tmp_path)))
    documento.save(original)

    quebrado = _corromper_imagem(original, tmp_path / "quebrado.docx")

    # Sem recuperação, a leitura falharia.
    with pytest.raises(zipfile.BadZipFile):
        docx.Document(str(quebrado))

    # Com recuperação, o texto continua acessível e a ficha é classificada.
    lido = ler_documento(str(quebrado))
    assert classificar(lido) is INDICADOR
    assert Extrator(INDICADOR).extrair(lido).valores["Eixo"] == "EIXO DE TESTE"


def _png_de_teste(pasta: Path) -> Path:
    from extrator.documento import _PNG_MINIMO

    caminho = pasta / "figura.png"
    caminho.write_bytes(_PNG_MINIMO)
    return caminho


# ---------------------------------------------------------------------------
# Variantes reais do modelo: "Bloco N:", rótulos explicados entre parênteses,
# "Financeiros" no lugar de "Orçamentárias/Financeiras" e a Ficha de Controle
# ---------------------------------------------------------------------------
CONTROLE = MODELOS_POR_CODIGO["CONTROLE"]


def criar_ficha_indicador_variante(destino: Path) -> Path:
    """Layout do eixo E13-GG: seções numeradas e rótulos com explicação."""
    documento = docx.Document()
    _tabela_vertical(documento, ["Bloco 1: VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_vertical(
        documento,
        [
            "Eixo",
            "Gestão Governamental",
            "Programa",
            "Governo Digital",
            "Compromisso (Objetivo do Compromisso)",
            "Aprimorar a governança de TIC",
            "Problema(s) vinculado(s) ao Compromisso",
            "Governança fragmentada",
            "Causa(s) Crítica(s)",
            "Causa X.",
            "Bloco 2: ATRIBUTOSDO INDICADOR DE COMPROMISSO",
        ],
    )
    _tabela_vertical(
        documento,
        [
            "Descrição (Descrição do Indicador)",
            "Percentual de implementação",
            "Fonte (Fonte de Informação)",
            "Saeb/SGI",
        ],
    )
    _tabela_vertical(documento, ["Bloco 4: INFORMAÇÕES COMPLEMENTARES"])
    _tabela_horizontal(
        documento,
        [
            ("Limitações para definição do valor da meta", ""),
            ("Operacionais", "Não identificado"),
            ("Financeiros", "Sem orçamento previsto"),
            ("Institucionais ou políticos", "Não identificado"),
        ],
    )
    documento.save(destino)
    return destino


def criar_ficha_controle(destino: Path) -> Path:
    """Capa do diretório do compromisso: rótulo e valor na mesma célula."""
    documento = docx.Document()
    tabela = documento.add_table(rows=6, cols=2)
    tabela.cell(0, 0).text = "NOME DO DIRETÓRIO DO COMPROMISSO: \nE13-GG-PGovernoDigital-C1"
    tabela.cell(1, 0).text = "EIXO: \n13 - Gestão Governamental"
    tabela.cell(2, 0).text = "PROGRAMA: \nGoverno Digital"
    tabela.cell(3, 0).text = "COMPROMISSO \nAprimorar a governança de TIC"
    tabela.cell(4, 0).text = "NOME DIGITADORA FIPLAN: \nManuela Alves"
    tabela.cell(4, 1).text = "DATA INSERÇÃO NO FIPLAN:\n11/ agosto / 2023"
    tabela.cell(5, 0).text = "N° de INDICADORES DE COMPROMISSO:\n3"
    tabela.cell(5, 1).text = "N° de Fichas de INICIATIVAS: \n2"
    documento.save(destino)
    return destino


def test_variante_com_bloco_numerado_e_reconhecida(tmp_path: Path):
    caminho = criar_ficha_indicador_variante(tmp_path / "variante.docx")
    documento = ler_documento(str(caminho))
    assert classificar(documento) is INDICADOR  # "Bloco 1: VÍNCULO..."


def test_variante_extrai_rotulos_com_explicacao(tmp_path: Path):
    caminho = criar_ficha_indicador_variante(tmp_path / "variante.docx")
    valores = Extrator(INDICADOR).extrair(ler_documento(str(caminho))).valores

    assert valores["Compromisso"] == "Aprimorar a governança de TIC"
    assert valores["Atributos_Descricao"] == "Percentual de implementação"
    assert valores["Fonte"] == "Saeb/SGI"
    # "Financeiros" no lugar de "Orçamentárias/Financeiras".
    assert valores["Limitacoes_Meta_Orcamentarias"] == "Sem orçamento previsto"
    # A causa crítica não invade o cabeçalho do bloco seguinte.
    assert valores["Causas_Criticas"] == "Causa X."


def test_ficha_de_controle_e_um_terceiro_tipo(tmp_path: Path):
    caminho = criar_ficha_controle(tmp_path / "controle.docx")
    documento = ler_documento(str(caminho))
    assert classificar(documento) is CONTROLE

    valores = Extrator(CONTROLE).extrair(documento).valores
    assert valores["Nome_Diretorio_Compromisso"] == "E13-GG-PGovernoDigital-C1"
    assert valores["Eixo"] == "13 - Gestão Governamental"
    # Rótulo na primeira linha e valor abaixo, sem dois-pontos.
    assert valores["Compromisso"] == "Aprimorar a governança de TIC"
    assert valores["Nome_Digitador_Fiplan"] == "Manuela Alves"
    assert valores["Data_Insercao_Fiplan"] == "11/ agosto / 2023"
    assert valores["Qtd_Indicadores_Compromisso"] == "3"
    assert valores["Qtd_Fichas_Iniciativas"] == "2"


def test_lote_com_os_tres_tipos(tmp_path: Path):
    pasta = tmp_path / "tres"
    pasta.mkdir()
    criar_ficha_indicador(pasta / "indicador.docx")
    criar_ficha_iniciativa(pasta / "iniciativa.docx")
    criar_ficha_controle(pasta / "controle.docx")

    registros, estatisticas = processar_lote(listar_documentos(pasta), pasta)
    assert estatisticas.processados == {"INDICADOR": 1, "INICIATIVA": 1, "CONTROLE": 1}
    assert estatisticas.total_ignorados == 0
    assert len(registros["CONTROLE"]) == 1


def criar_iniciativa_variante_teto(destino: Path) -> Path:
    """Layout do eixo E10-MASH: "Previsão de Recursos" e "Total do Teto"."""
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(documento, ["Eixo", "MEIO AMBIENTE"])
    _tabela_vertical(documento, ["ATRIBUTOS DA INICIATIVA"])
    _tabela_grade(
        documento,
        [
            ["Fatores Críticos de Contexto da Iniciativa", ""],
            ["Operacionais", "Equipe reduzida"],
            ["Financeiros", "Alto custo da contratação"],
        ],
    )
    _tabela_grade(
        documento,
        [
            ["Previsão de Recursos – Estimativa do Teto da Iniciativa", "", ""],
            ["Código da Fonte", "Nome da Fonte", "Montante em R$"],
            ["100", "Tesouro Estadual", "R$ 5.000,00"],
            ["Total do Teto", "", "R$ 5.000,00"],
        ],
    )
    documento.save(destino)
    return destino


def test_variante_previsao_de_recursos_com_total_do_teto(tmp_path: Path):
    caminho = criar_iniciativa_variante_teto(tmp_path / "teto.docx")
    valores = Extrator(INICIATIVA).extrair(ler_documento(str(caminho))).valores

    linhas = valores["Recursos_Orcamentarios"].split("\n")
    assert linhas[0] == (
        "Código da Fonte: 100 | Nome da Fonte: Tesouro Estadual | "
        "Montante em R$: R$ 5.000,00"
    )
    # O rótulo do total é preservado como está no documento.
    assert linhas[-1] == "Total do Teto: R$ 5.000,00"
    assert valores["Fatores_Criticos_Orcamentarios"] == "Alto custo da contratação"


def test_variante_previsao_de_recursos_e_localizada(tmp_path: Path):
    caminho = criar_iniciativa_variante_teto(tmp_path / "teto.docx")
    resultado = Extrator(INICIATIVA).extrair(ler_documento(str(caminho)))
    assert "Recursos_Orcamentarios" not in resultado.nao_encontrados


def criar_indicador_vinculo_do_compromisso(destino: Path) -> Path:
    """Layout do E11-SAÚDE: "INDICADOR DO COMPROMISSO" e duas colunas.

    Aqui "Problema(s) vinculado(s)" e "Causas Críticas" ficam lado a lado, cada
    um com sua própria coluna de itens.
    """
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DO COMPROMISSO"])
    _tabela_grade(
        documento,
        [
            ["Eixo", ""],
            ["Saúde", ""],
            ["Problema(s) vinculado(s) ao Compromisso", "Causas Críticas"],
            ["Primeiro problema", "Primeira causa"],
            ["Segundo problema", "Segunda causa"],
        ],
    )
    _tabela_vertical(
        documento,
        ["ATRIBUTOS DO INDICADOR DO COMPROMISSO"],
    )
    _tabela_vertical(documento, ["Descrição", "Número de pesquisas incorporadas"])
    documento.save(destino)
    return destino


def test_variante_indicador_do_compromisso(tmp_path: Path):
    caminho = criar_indicador_vinculo_do_compromisso(tmp_path / "saude.docx")
    documento = ler_documento(str(caminho))
    assert classificar(documento) is INDICADOR  # "DO" no lugar de "DE"

    valores = Extrator(INDICADOR).extrair(documento).valores
    assert valores["Eixo"] == "Saúde"
    assert valores["Atributos_Descricao"] == "Número de pesquisas incorporadas"
    # Colunas lado a lado: cada bloco desce pela sua própria coluna.
    assert valores["Problemas_Vinculados"] == "Primeiro problema\nSegundo problema"
    assert valores["Causas_Criticas"] == "Primeira causa\nSegunda causa"


def test_variante_do_compromisso_gera_uma_linha_por_problema(tmp_path: Path):
    pasta = tmp_path / "saude"
    pasta.mkdir()
    criar_indicador_vinculo_do_compromisso(pasta / "saude.docx")

    registros, estatisticas = processar_lote(listar_documentos(pasta), pasta)
    assert estatisticas.processados["INDICADOR"] == 1
    assert estatisticas.linhas["INDICADOR"] == 2
    assert [l["Problemas_Vinculados"] for l in registros["INDICADOR"]] == [
        "Primeiro problema",
        "Segundo problema",
    ]


def test_ficha_de_controle_redacoes_da_contagem(tmp_path: Path):
    """As três redações do rótulo de contagem que existem no acervo."""
    documento = docx.Document()
    tabela = documento.add_table(rows=4, cols=1)
    tabela.cell(0, 0).text = "NOME DO DIRETÓRIO DO COMPROMISSO (copiar P:\\Fichas): \nE10-MASH-C1"
    # Com "Fichas de" no meio e a lista de códigos abaixo da contagem.
    tabela.cell(1, 0).text = (
        "Número de Fichas de INDICADORES DE COMPROMISSO:\n03 \ncód 1646\nCód 1647\n"
    )
    tabela.cell(2, 0).text = "Número de Fichas de INICIATIVAS: \n04\nCód 1646 IN 0001"
    tabela.cell(3, 0).text = "PENDÊNCIAS E OBSERVAÇÕES \nFinalizado"
    caminho = tmp_path / "controle_variante.docx"
    documento.save(caminho)

    valores = Extrator(CONTROLE).extrair(ler_documento(str(caminho))).valores
    # Só a contagem entra; a lista de códigos abaixo não faz parte do valor.
    assert valores["Qtd_Indicadores_Compromisso"] == "03"
    assert valores["Qtd_Fichas_Iniciativas"] == "04"
    assert valores["Nome_Diretorio_Compromisso"] == "E10-MASH-C1"
    assert valores["Pendencias_Observacoes"] == "Finalizado"


def test_ficha_de_controle_contagem_sem_dois_pontos(tmp_path: Path):
    documento = docx.Document()
    tabela = documento.add_table(rows=2, cols=2)
    tabela.cell(0, 0).text = "NOME DO DIRETÓRIO DO COMPROMISSO: \nE1-ASGD-C9"
    tabela.cell(1, 0).text = "N° de Fichas de INDICADORES \n 2\nE1- ASGD Indicador 1"
    tabela.cell(1, 1).text = "N° de Fichas de INICIATIVAS: \n2"
    caminho = tmp_path / "controle_sem_dois_pontos.docx"
    documento.save(caminho)

    valores = Extrator(CONTROLE).extrair(ler_documento(str(caminho))).valores
    assert valores["Qtd_Indicadores_Compromisso"] == "2"
    assert valores["Qtd_Fichas_Iniciativas"] == "2"


# ---------------------------------------------------------------------------
# Erros de captura do rótulo "Problema(s) vinculado(s) ao Compromisso"
#
# Os três casos vieram da conferência do acervo real: colunas de problema em
# branco (com Status OK, sem aviso nenhum) e colunas preenchidas com sujeira.
# ---------------------------------------------------------------------------
def _ficha_com_quadro_de_orientacoes(destino: Path) -> Path:
    """Ficha em que o rótulo aparece duas vezes na mesma seção.

    A primeira ocorrência é o quadro de orientações de preenchimento, sem
    resposta ao lado; a ficha preenchida vem logo depois.
    """
    documento = docx.Document()

    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_horizontal(
        documento,
        [
            ("Problema(s) vinculado(s) ao Compromisso", ""),
            ("Causa(s) Crítica(s)", ""),
        ],
    )
    _tabela_horizontal(
        documento,
        [
            ("Eixo", "Segurança Pública"),
            ("Compromisso", "Fortalecer a Polícia Comunitária"),
            ("Problema(s) vinculado(s) ao Compromisso", "P1 Violência urbana"),
            ("Causa(s) Crítica(s)", "CC1 Baixa cobertura"),
        ],
    )
    documento.save(destino)
    return destino


def test_rotulo_do_quadro_de_orientacoes_nao_esvazia_a_coluna(tmp_path: Path):
    """A busca segue até a ocorrência que tem resposta."""
    caminho = _ficha_com_quadro_de_orientacoes(tmp_path / "orientacoes.docx")
    resultado = Extrator(INDICADOR).extrair(ler_documento(str(caminho)))

    assert resultado.valores["Problemas_Vinculados"] == "P1 Violência urbana"
    assert resultado.valores["Causas_Criticas"] == "CC1 Baixa cobertura"
    assert "Problemas_Vinculados" not in resultado.nao_encontrados


def test_caixa_de_selecao_nao_vira_problema(tmp_path: Path):
    """Célula com um caractere solto ("e" de fonte de símbolos) não é valor."""
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_horizontal(
        documento,
        [
            ("Compromisso", "Compromisso qualquer"),
            ("Problema(s) vinculado(s) ao Compromisso", "e"),
            ("Causa(s) Crítica(s)", "."),
        ],
    )
    caminho = tmp_path / "sujeira.docx"
    documento.save(caminho)

    resultado = Extrator(INDICADOR).extrair(ler_documento(str(caminho)))
    assert resultado.valores["Problemas_Vinculados"] == ""
    assert resultado.valores["Causas_Criticas"] == ""
    # E o arquivo é apontado no relatório, em vez de sair calado.
    assert "Problemas_Vinculados" in resultado.nao_encontrados
    assert "Causas_Criticas" in resultado.nao_encontrados


def test_rotulo_sem_resposta_entra_em_campos_nao_encontrados(tmp_path: Path):
    """Rótulo presente e célula vazia: a ficha precisa aparecer como pendência."""
    pasta = tmp_path / "sem_resposta"
    pasta.mkdir()
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DO INDICADOR DE COMPROMISSO"])
    _tabela_horizontal(
        documento,
        [
            ("Compromisso", "Compromisso sem problema preenchido"),
            ("Problema(s) vinculado(s) ao Compromisso", ""),
        ],
    )
    documento.save(pasta / "ficha.docx")

    registros, estatisticas = processar_lote(listar_documentos(pasta), pasta)
    linha = registros["INDICADOR"][0]
    assert linha["Problemas_Vinculados"] == ""
    assert linha["Status"] == "OK_COM_PENDENCIAS"
    assert "Problemas_Vinculados" in linha["Campos_Nao_Encontrados"]
    assert estatisticas.campos_ausentes["INDICADOR"]["Problemas_Vinculados"] == 1


def test_valor_curto_legitimo_continua_valendo(tmp_path: Path):
    """A regra de sujeira não pode derrubar um valor curto de verdade."""
    documento = docx.Document()
    _tabela_vertical(documento, ["ATRIBUTOS DO INDICADOR DE COMPROMISSO"])
    _tabela_horizontal(
        documento,
        [
            ("Unidade de medida", "%"),
            ("Valor da meta", "5"),
            ("Sigla do Órgão", "SSP"),
        ],
    )
    caminho = tmp_path / "curtos.docx"
    documento.save(caminho)

    valores = Extrator(INDICADOR).extrair(ler_documento(str(caminho))).valores
    assert valores["Valor_Meta"] == "5"
    assert valores["Responsavel_Sigla_Orgao"] == "SSP"


# ---------------------------------------------------------------------------
# Layouts conferidos contra as fichas reais do acervo
# ---------------------------------------------------------------------------
def _ficha_com_rotulo_apagado(destino: Path) -> Path:
    """Ficha real em que o rótulo "Causa(s) Crítica(s)" foi apagado.

    Sobrou a célula em branco no lugar do rótulo, com as causas logo abaixo.
    Reproduz E12-SP-PSegPubDefesaSocial- C4 (CorregedoriaCBM).
    """
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Compromisso",
            "Fortalecer as ações das corregedorias",
            "Problema(s) vinculado(s) ao Compromisso",
            "A corrupção de agentes públicos.",
            "",  # aqui estava o rótulo "Causa(s) Crítica(s)"
            "C4P1CC1: Baixa divulgação dos canais de denúncia.\nC4P1CC2: Ausência de sistema.",
            "Ação(ões) Crítica(s)",
            "AC1: Levantar os requisitos comuns.",
        ],
    )
    documento.save(destino)
    return destino


def test_celula_vazia_encerra_a_lista(tmp_path: Path):
    """Sem o rótulo seguinte, a lista parava só no bloco de baixo e o engolia."""
    caminho = _ficha_com_rotulo_apagado(tmp_path / "rotulo_apagado.docx")
    resultado = Extrator(INICIATIVA).extrair(ler_documento(str(caminho)))

    assert resultado.valores["Problemas_Vinculados"] == "A corrupção de agentes públicos."
    assert resultado.valores["Acoes_Criticas"] == "AC1: Levantar os requisitos comuns."
    # A causa não tem rótulo no documento: fica vazia e é apontada, não some
    # dentro da coluna de problemas.
    assert resultado.valores["Causas_Criticas"] == ""
    assert "Causas_Criticas" in resultado.nao_encontrados


def test_anotacao_de_trabalho_vem_como_esta_na_ficha(tmp_path: Path):
    """Anotação de quem preencheu ("AC 4,5,7,8") no fim da lista de problemas.

    Reproduz E6-DU-PMelhoriaCond-C2. É texto digitado na ficha: o aplicativo
    traz como está, e a limpeza é feita no documento. Só o que não tem
    conteúdo algum (uma letra solta, um ponto) é separação entre itens.
    """
    pasta = tmp_path / "anotacao"
    pasta.mkdir()
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Problema(s) vinculado(s) ao Compromisso",
            "P1-Inadequadas condições de acessibilidade\n"
            "P2-Dificuldades nas condições de mobilidade\n"
            "AC 4,5,7,810,12,13,14",
        ],
    )
    documento.save(pasta / "ficha.docx")

    registros, _ = processar_lote(listar_documentos(pasta), pasta)
    assert [linha["Problemas_Vinculados"] for linha in registros["INICIATIVA"]] == [
        "P1-Inadequadas condições de acessibilidade",
        "P2-Dificuldades nas condições de mobilidade",
        "AC 4,5,7,810,12,13,14",
    ]


def test_codigo_sozinho_na_celula_continua_valendo(tmp_path: Path):
    """Célula que só tem "P1" é o que está na ficha, e vai para a planilha."""
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento, ["Problema(s) vinculado(s) ao Compromisso", "P1", "Ação(ões) Crítica(s)"]
    )
    caminho = tmp_path / "codigo.docx"
    documento.save(caminho)

    resultado = Extrator(INICIATIVA).extrair(ler_documento(str(caminho)))
    assert resultado.valores["Problemas_Vinculados"] == "P1"
    assert "Problemas_Vinculados" not in resultado.nao_encontrados


def test_ligacao_solta_nao_mexe_em_campo_de_codigo(tmp_path: Path):
    """A regra derruba letra solta e pontuação; sigla e código continuam."""
    from extrator.parser import sem_ligacoes_soltas

    assert sem_ligacoes_soltas("P1 Violência\ne\nP2 Racismo") == "P1 Violência\nP2 Racismo"
    assert sem_ligacoes_soltas("P1 Violência\n.\nP2 Racismo") == "P1 Violência\nP2 Racismo"
    # Texto com conteúdo, por mais curto ou codificado que seja, permanece.
    for texto in ("AC 4,5,7,810,12,13,14", "C5P1,3CC12,13,16AC3", "P1", "Sedentarismo"):
        assert sem_ligacoes_soltas(texto) == texto

    documento = docx.Document()
    _tabela_vertical(documento, ["ATRIBUTOS DA INICIATIVA"])
    _tabela_grade(documento, [["Sigla do Órgão", "UO"], ["SSP", "20803"]])
    caminho = tmp_path / "codigos.docx"
    documento.save(caminho)

    valores = Extrator(INICIATIVA).extrair(ler_documento(str(caminho))).valores
    assert valores["Responsavel_Sigla_Orgao"] == "SSP"
    assert valores["Responsavel_UO"] == "20803"


def test_ligacao_e_ponto_solto_separam_itens_em_vez_de_virar_item(tmp_path: Path):
    """Lista escrita com "e" de ligação e um ponto de digitação no meio.

    Nas fichas a lista às vezes vem "P1 ... / e / P2 ...", com o "e" sozinho na
    linha, e às vezes sobra um "." solto de digitação. Os dois são separação
    entre itens, não itens: não podem virar uma linha da planilha.
    """
    pasta = tmp_path / "ligacao"
    pasta.mkdir()
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Problema(s) vinculado(s) ao Compromisso",
            "P1 Violência contra jovens negros\ne\nP2 Racismo institucional\n.\n"
            "P3 Baixa escolaridade",
        ],
    )
    documento.save(pasta / "ficha.docx")

    registros, _ = processar_lote(listar_documentos(pasta), pasta)
    assert [linha["Problemas_Vinculados"] for linha in registros["INICIATIVA"]] == [
        "P1 Violência contra jovens negros",
        "P2 Racismo institucional",
        "P3 Baixa escolaridade",
    ]


def test_ponto_final_dentro_da_frase_nao_separa(tmp_path: Path):
    """O ponto só é descartado quando está sozinho: frase não é picotada."""
    pasta = tmp_path / "frase"
    pasta.mkdir()
    documento = docx.Document()
    _tabela_vertical(documento, ["VÍNCULO DA INICIATIVA"])
    _tabela_vertical(
        documento,
        [
            "Problema(s) vinculado(s) ao Compromisso",
            "P1 Baixa cobertura. A rede não alcança o interior. Faltam equipes.",
        ],
    )
    documento.save(pasta / "ficha.docx")

    registros, _ = processar_lote(listar_documentos(pasta), pasta)
    linhas = registros["INICIATIVA"]
    assert len(linhas) == 1
    assert linhas[0]["Problemas_Vinculados"] == (
        "P1 Baixa cobertura. A rede não alcança o interior. Faltam equipes."
    )
